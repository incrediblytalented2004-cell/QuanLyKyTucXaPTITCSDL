import sys
import subprocess
import os

try:
    import docx
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = Document()


heading = doc.add_heading('BÁO CÁO YÊU CẦU LƯU TRỮ VÀ MÔ HÌNH RÀNG BUỘC', 0)
heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_heading('Đề tài: Quản lý dữ liệu Ký túc xá Sinh viên', level=1)


doc.add_heading('1. Các dữ liệu cần lưu trữ', level=2)
doc.add_paragraph('Hệ thống cần thu thập và phân mảnh các đối tượng lưu trữ cơ bản thay vì để lộn xộn vào một nơi:')

doc.add_paragraph('Tòa nhà: Lưu trữ định danh Tòa nhà, Tên tòa nhà.', style='List Bullet')
doc.add_paragraph('Sinh viên: Lưu trữ Tên sinh viên, Năm sinh, Mã số định danh.', style='List Bullet')
doc.add_paragraph('Phòng: Lưu trữ thông tin Phòng, Tình trạng phòng, Số người đang lưu trú.', style='List Bullet')
doc.add_paragraph('Hợp đồng: Lưu trữ chi tiết Hợp đồng cấp quyền lưu trú.', style='List Bullet')


doc.add_heading('2. Mối quan hệ liên kết giữa các thực thể', level=2)
doc.add_paragraph('Khối dữ liệu không hề rời rạc độc lập (đặc biệt là Sinh Viên) mà bị ràng buộc chặt chẽ với nhau thông qua mối quan hệ Khoá chính - Khoá ngoại:')

p_quanhe1 = doc.add_paragraph(style='List Bullet')
p_quanhe1.add_run('Mối quan hệ Tòa nhà - Phòng (1:N):').bold = True
p_quanhe1.add_run(' Một tòa nhà sở hữu nhiều phòng. Sự ràng buộc: Không thể có một phòng tồn tại trôi nổi nếu không nằm trong một tòa nhà có thật.')

p_quanhe2 = doc.add_paragraph(style='List Bullet')
p_quanhe2.add_run('Mối quan hệ của Sinh viên (Ràng buộc gốc của Hợp đồng):').bold = True
p_quanhe2.add_run(' Sinh viên không hề nằm độc lập. Sinh Viên chính là đối tượng lõi tương tác với Phòng thông qua "Mạng lưới Hợp đồng". Mối liên hệ đa nhãn này cho thấy: Một Hợp đồng bắt buộc phải chỉ định từ Tên/Mã của 1 Sinh Viên đang thực sự có tồn tại trên hệ thống và gắn với 1 Phòng có thật. Do đó sự ràng buộc của Sinh viên là: Gắn kết trực tiếp tính pháp lý của Hợp đồng.')


doc.add_heading('3. Các Ràng buộc (Constraints) áp đặt lên hệ thống', level=2)

p_rb1 = doc.add_paragraph(style='List Bullet')
p_rb1.add_run('Tính sống còn của thông tin Sinh viên:').bold = True
p_rb1.add_run(' Ràng buộc dữ liệu chống mồ côi (No Orphan Records). Khi Sinh viên đang có một Hợp đồng thuê phòng chưa hết hạn, quản trị viên KHÔNG THỂ tùy tiện xoá dữ liệu Sinh viên đó hoặc xóa Phòng đó. Nó tạo sự ràng buộc toàn vẹn bảo vệ Sinh viên.')

p_rb2 = doc.add_paragraph(style='List Bullet')
p_rb2.add_run('Ràng buộc chống vượt rào của Hợp đồng (Trigger):').bold = True
p_rb2.add_run(' Trước khi tạo ra Hợp đồng gắn Sinh viên vào Phòng, thủ tục quản lý sẽ rà soát [Số người] hiện hành trong phòng. Nếu đã vượt ngưỡng tiêu chuẩn, Ràng buộc chối từ Hợp đồng sẽ kích hoạt.')

p_rb3 = doc.add_paragraph(style='List Bullet')
p_rb3.add_run('Ràng buộc chuẩn hóa dữ liệu Sinh viên:').bold = True
p_rb3.add_run(' Họ tên sinh viên, Năm sinh không được phép để trống (Not Null constraints). Có thể yêu cầu kiểm tra ràng tuổi của Sinh viên thông qua độ tuổi từ Năm sinh.')

doc.save(os.path.join(os.getcwd(), 'Bao_Cao_Yeu_Cau_Va_Moi_Quan_He_Rang_Buoc.docx'))
print("Complete")
